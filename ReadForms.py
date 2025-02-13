import docx
import os
import pandas as pd
import re

def listar(directorio):
  excels = []
  for archivo in os.listdir(directorio):
    if archivo.endswith(".docx"):
      excels.append(directorio+archivo)
  return excels

#Write your own path
ruta = ""

docs = listar(ruta)

df = pd.DataFrame(columns=['Usuario','Equipo', 'S/N', 'Fecha entrega', 'Observaciones', 'Fecha devolución', 'Estado devolucion', 'Firma devolución'])
for word in docs:
  doc = docx.Document(word)
  tables = doc.tables
  paragraphs = doc.paragraphs
  pattern = re.compile("Name:")
  for par in paragraphs:
    match = pattern.match(par.text)
    if match:
     name = par.text[5:]
  try:
    tabla = tables[0]
    lista=[]
    for row in tabla.rows:
        lista.append([ [name] + [cell.text for cell in row.cells]])
    clean = [[elemento for sublista in fila for elemento in sublista] for fila in lista]
    clean.pop(0)
    df= pd.concat([df, pd.DataFrame(clean)])
  except:
    print(word)

df.to_excel(ruta+"resumen.xlsx",index=False)
